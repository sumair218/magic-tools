import os
import uuid
from django.conf import settings
from django.contrib import messages
from django.http import FileResponse, HttpResponseRedirect
from django.shortcuts import redirect, render
from django.urls import reverse

from .forms import CompressPDFForm, MergePDFForm, PDFUploadForm, SplitPDFForm, WordUploadForm

from pypdf import PdfWriter, PdfReader
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document


def home(request):
    return render(request, 'pdf_tools/home.html')


def _save_temp_file(uploaded_file):
    temp_dir = settings.MEDIA_ROOT / 'temp'
    temp_dir.mkdir(parents=True, exist_ok=True)
    filename = f"{uuid.uuid4().hex}_{uploaded_file.name}"
    file_path = temp_dir / filename
    with open(file_path, 'wb+') as destination:
        for chunk in uploaded_file.chunks():
            destination.write(chunk)
    return file_path


def _serve_file(path, filename):
    response = FileResponse(open(path, 'rb'), as_attachment=True, filename=filename)
    return response


def pdf_to_word(request):
    form = PDFUploadForm(request.POST or None, request.FILES or None)
    download_url = None
    download_name = None

    if request.method == 'POST' and form.is_valid():
        pdf_path = _save_temp_file(form.cleaned_data['pdf_file'])
        doc_path = settings.MEDIA_ROOT / 'temp' / f"{uuid.uuid4().hex}.docx"

        reader = PdfReader(pdf_path)
        doc = Document()

        for page in reader.pages:
            text = page.extract_text() or ''
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.add_page_break()

        doc.save(doc_path)
        return _serve_file(doc_path, f"{os.path.splitext(form.cleaned_data['pdf_file'].name)[0]}.docx")

    return render(request, 'pdf_tools/tool_form.html', {
        'tool_name': 'PDF to Word',
        'tool_description': 'Upload a PDF file and convert it to an editable Word document.',
        'form': form,
        'download_url': download_url,
        'download_name': download_name,
    })


def word_to_pdf(request):
    form = WordUploadForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        doc_path = _save_temp_file(form.cleaned_data['doc_file'])
        pdf_path = settings.MEDIA_ROOT / 'temp' / f"{uuid.uuid4().hex}.pdf"

        doc = Document(doc_path)
        c = canvas.Canvas(str(pdf_path), pagesize=letter)

        for para in doc.paragraphs:
            c.drawString(72, 720, para.text)
            c.showPage()

        c.save()
        return _serve_file(pdf_path, f"{os.path.splitext(form.cleaned_data['doc_file'].name)[0]}.pdf")

    return render(request, 'pdf_tools/tool_form.html', {
        'tool_name': 'Word to PDF',
        'tool_description': 'Upload a Word document and convert it into a PDF file.',
        'form': form,
    })


def merge_pdf(request):
    form = MergePDFForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        files = request.FILES.getlist('pdf_files')
        pdf_writer = PdfWriter()

        for uploaded_file in files:
            pdf_path = _save_temp_file(uploaded_file)
            pdf_reader = PdfReader(pdf_path)
            for page in pdf_reader.pages:
                pdf_writer.add_page(page)

        output_path = settings.MEDIA_ROOT / 'temp' / f"{uuid.uuid4().hex}.pdf"
        with open(output_path, 'wb') as out_file:
            pdf_writer.write(out_file)

        return _serve_file(output_path, 'merged.pdf')

    return render(request, 'pdf_tools/tool_form.html', {
        'tool_name': 'Merge PDF',
        'tool_description': 'Combine multiple PDF files into one merged document.',
        'form': form,
    })


def split_pdf(request):
    form = SplitPDFForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        pdf_path = _save_temp_file(form.cleaned_data['pdf_file'])
        reader = PdfReader(pdf_path)
        start_page = form.cleaned_data.get('start_page') or 1
        end_page = form.cleaned_data.get('end_page') or len(reader.pages)

        output_path = settings.MEDIA_ROOT / 'temp' / f"{uuid.uuid4().hex}.pdf"
        writer = PdfWriter()

        for i in range(start_page - 1, min(end_page, len(reader.pages))):
            writer.add_page(reader.pages[i])

        with open(output_path, 'wb') as out_file:
            writer.write(out_file)

        return _serve_file(output_path, f"split_{start_page}_to_{end_page}.pdf")

    return render(request, 'pdf_tools/tool_form.html', {
        'tool_name': 'Split PDF',
        'tool_description': 'Split a PDF into separate pages or smaller files.',
        'form': form,
    })


def compress_pdf(request):
    form = CompressPDFForm(request.POST or None, request.FILES or None)
    if request.method == 'POST' and form.is_valid():
        pdf_path = _save_temp_file(form.cleaned_data['pdf_file'])
        reader = PdfReader(pdf_path)
        writer = PdfWriter()

        for page in reader.pages:
            writer.add_page(page)

        output_path = settings.MEDIA_ROOT / 'temp' / f"{uuid.uuid4().hex}.pdf"
        with open(output_path, 'wb') as out_file:
            writer.write(out_file)

        return _serve_file(output_path, f"compressed_{form.cleaned_data['pdf_file'].name}")

    return render(request, 'pdf_tools/tool_form.html', {
        'tool_name': 'Compress PDF',
        'tool_description': 'Reduce the file size of a PDF while keeping the content readable.',
        'form': form,
    })
